"""
Update the template .docx with actual DeliveryIQ coursework content.
Preserves the first 3 tables (title page, assignment, calendar plan).
Replaces all body content from Реферат onward.
Formatting: Times New Roman 14pt, 1.5 line spacing, 1.25cm first-line indent, justify.
"""

import copy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

BASE = Path(__file__).parent
DIAGRAMS = BASE / "diagrams"
TEMPLATE = BASE / "шаблон_КР_ПІС.docx"
OUTPUT = BASE / "шаблон_КР_ПІС_filled.docx"

# ── Formatting helpers ──────────────────────────────────────────────────────

def fmt_run(run, size=Pt(14), bold=None, italic=False, font_name='Times New Roman', color=None):
    run.font.name = font_name
    run.font.size = size
    if bold is not None:
        run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # East-Asian font fallback
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
        rpr.insert(0, rfonts)
    else:
        rfonts.set(qn('w:ascii'), font_name)
        rfonts.set(qn('w:hAnsi'), font_name)
        rfonts.set(qn('w:cs'), font_name)


def fmt_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=Cm(1.25),
                  space_before=Pt(0), space_after=Pt(0), line_spacing=1.5,
                  keep_next=False):
    pf = p.paragraph_format
    pf.alignment = alignment
    pf.first_line_indent = first_indent
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    if keep_next:
        pf.keep_with_next = True


def add_body_paragraph(doc, text, bold=None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       first_indent=Cm(1.25), font_name='Times New Roman',
                       size=Pt(14), space_after=Pt(0)):
    p = doc.add_paragraph()
    fmt_paragraph(p, alignment=alignment, first_indent=first_indent, space_after=space_after)
    run = p.add_run(text)
    fmt_run(run, size=size, bold=bold, font_name=font_name)
    return p


def add_heading_centered(doc, text, level=1):
    """Add a section heading: centered, bold, uppercase, no indent."""
    p = doc.add_paragraph()
    fmt_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Cm(0),
                  space_before=Pt(0), space_after=Pt(12))
    run = p.add_run(text)
    fmt_run(run, size=Pt(14), bold=True)
    return p


def add_subheading(doc, text):
    """Add a subsection heading: left-aligned with indent, bold."""
    p = doc.add_paragraph()
    fmt_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=Cm(1.25),
                  space_before=Pt(12), space_after=Pt(6))
    run = p.add_run(text)
    fmt_run(run, size=Pt(14), bold=True)
    return p


def add_bullet_item(doc, text, bold_prefix=None):
    """Add a dash-prefixed list item."""
    p = doc.add_paragraph()
    fmt_paragraph(p, first_indent=Cm(1.25))
    if bold_prefix:
        r = p.add_run(f'– {bold_prefix}')
        fmt_run(r, bold=True)
        r2 = p.add_run(f' – {text}')
        fmt_run(r2)
    else:
        r = p.add_run(f'– {text}')
        fmt_run(r)
    return p


def add_numbered_item(doc, number, text):
    p = doc.add_paragraph()
    fmt_paragraph(p, first_indent=Cm(1.25))
    run = p.add_run(f'{number}) {text}')
    fmt_run(run)
    return p


def add_image_with_caption(doc, image_name, caption, width=Cm(16)):
    """Add centered image with italic caption below."""
    img_path = DIAGRAMS / image_name
    if not img_path.exists():
        add_body_paragraph(doc, f'[Зображення {image_name} не знайдено]',
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Cm(0))
        return

    p = doc.add_paragraph()
    fmt_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Cm(0))
    run = p.add_run()
    run.add_picture(str(img_path), width=width)

    p2 = doc.add_paragraph()
    fmt_paragraph(p2, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Cm(0),
                  space_after=Pt(12))
    run2 = p2.add_run(caption)
    fmt_run(run2, size=Pt(12), italic=True)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run._element.append(parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>'))


# ── Remove all body elements after the front-matter tables ──────────────────

def strip_body_after_tables(doc):
    """Remove all body elements starting from the Реферат paragraph (Body[6] onward),
    keeping the 3 tables and their spacing paragraphs."""
    body = doc.element.body
    children = list(body)

    # Find the index of Body[6] — first paragraph after 3rd table
    # Structure: tbl, p, tbl, p, tbl, p, p(Реферат), ...
    table_count = 0
    cut_index = None
    for i, child in enumerate(children):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'tbl':
            table_count += 1
        if table_count >= 3 and tag == 'p':
            # After the 3rd table, skip the spacer paragraph
            text = ''.join(child.itertext()).strip()
            if not text:
                continue
            # This is the first content paragraph after tables (Реферат)
            cut_index = i
            break

    if cut_index is None:
        raise RuntimeError("Could not find body content start after tables")

    # Keep sectPr (last child is typically section properties)
    sect_pr = None
    last = children[-1]
    if last.tag.endswith('}sectPr') or last.tag == 'sectPr':
        sect_pr = last

    # Remove everything from cut_index onward
    to_remove = children[cut_index:]
    for child in to_remove:
        body.remove(child)

    # Re-add sectPr if it was removed
    if sect_pr is not None:
        body.append(sect_pr)


# ── Update assignment page (table 2) topic ──────────────────────────────────

def update_assignment_table(doc):
    """Update the assignment table (2nd table) with DeliveryIQ topic and input data."""
    body = doc.element.body
    tables = [c for c in body if c.tag.endswith('}tbl') or ('}' not in c.tag and c.tag == 'tbl')]

    if len(tables) < 2:
        return

    tbl = tables[1]  # 2nd table = assignment page

    # Find and replace "Індивідуальна предметна область" with our topic
    for cell in tbl.iter(qn('w:t')):
        text = cell.text or ''
        if 'Індивідуальна предметна' in text:
            cell.text = text.replace('Індивідуальна предметна', 'оптимізації розподілу вантажів у транспортній мережі')
        if 'область' in text and 'предметної' not in text and 'наочної' not in text:
            cell.text = text.replace('область', '(DeliveryIQ)')

    # Find and replace input data placeholder
    for cell in tbl.iter(qn('w:t')):
        text = cell.text or ''
        if 'Необхідно описати' in text:
            cell.text = 'Адреси точок доставки, параметри транспортних'
        if 'даними для створювального' in text:
            cell.text = 'засобів (тип, місткість), дані вуличної мережі OpenStreetMap, координати депо.'


# ── Build content ───────────────────────────────────────────────────────────

def add_referat(doc, skip_page_break=False):
    """Add the Реферат (abstract) section."""
    if not skip_page_break:
        add_page_break(doc)

    p = doc.add_paragraph()
    fmt_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Cm(0),
                  space_after=Pt(12))
    run = p.add_run('Реферат')
    fmt_run(run, bold=True)

    doc.add_paragraph()  # spacer

    add_body_paragraph(
        doc,
        'Пояснювальна записка: __ с., 10 рис., __ джерел.'
    )
    doc.add_paragraph()

    add_body_paragraph(
        doc,
        'ПРОЄКТУВАННЯ, ІНФОРМАЦІЙНА СИСТЕМА, UML, ОПТИМІЗАЦІЯ МАРШРУТІВ, '
        'ЗАДАЧА КОМІВОЯЖЕРА, OPENSTREETMAP, МУЛЬТИМОДАЛЬНА МАРШРУТИЗАЦІЯ.',
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Cm(0)
    )
    doc.add_paragraph()

    add_body_paragraph(
        doc,
        'Об\'єктом курсової роботи є оптимізація маршрутів доставки вантажів '
        'у міській транспортній мережі.'
    )
    add_body_paragraph(
        doc,
        'Предмет роботи \u2013 інформаційна система оптимізації розподілу '
        'вантажів у транспортній мережі (DeliveryIQ).'
    )
    add_body_paragraph(
        doc,
        'Мета роботи \u2013 проєктування інформаційної системи оптимізації '
        'маршрутів доставки з використанням об\'єктно-орієнтованого підходу '
        'та мови моделювання UML.'
    )
    add_body_paragraph(
        doc,
        'Завдання роботи: аналіз предметної області маршрутизації транспортних '
        'засобів, проєктування інформаційної системи з використанням UML-діаграм, '
        'опис реалізації ключових алгоритмів оптимізації.'
    )
    add_body_paragraph(
        doc,
        'Методи дослідження \u2013 об\'єктно-орієнтоване моделювання.'
    )
    add_body_paragraph(
        doc,
        'Засоби розробки \u2013 мова програмування Python 3.11, фреймворк '
        'Streamlit, бібліотеки OSMnx та NetworkX, засіб моделювання PlantUML '
        'та мова моделювання UML.'
    )
    add_body_paragraph(
        doc,
        'Структура роботи включає три розділи. Загальний розділ містить аналіз '
        'предметної області, огляд існуючих рішень та постановку задачі. '
        'В спеціальному розділі представлено проєктування інформаційної системи '
        'з використанням UML-діаграм. В розрахунковому розділі описано вибір '
        'засобів розробки, структуру програмного забезпечення та ключові алгоритми.'
    )
    add_body_paragraph(
        doc,
        'Робота має практичне значення та може бути використана при створенні '
        'відповідної інформаційної системи оптимізації маршрутів доставки.'
    )


def add_toc(doc):
    """Add the table of contents."""
    add_page_break(doc)

    p = doc.add_paragraph()
    fmt_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=Cm(0),
                  space_after=Pt(12))
    run = p.add_run('ЗМІСТ')
    fmt_run(run, bold=True)

    doc.add_paragraph()

    items = [
        'ВСТУП',
        '1 ЗАГАЛЬНИЙ РОЗДІЛ',
        '  1.1 Опис предметної області',
        '  1.2 Огляд існуючих рішень',
        '  1.3 Постановка задачі',
        '2 СПЕЦІАЛЬНИЙ РОЗДІЛ',
        '  2.1 Діаграма прецедентів (Use Case Diagram)',
        '  2.2 Діаграми класів (Class Diagram)',
        '    2.2.1 Діаграма класів предметної області',
        '    2.2.2 Діаграма класів модульної структури',
        '  2.3 Діаграми послідовності (Sequence Diagram)',
        '    2.3.1 Процес оптимізації маршруту (TSP)',
        '    2.3.2 Геокодування адреси доставки',
        '    2.3.3 Оптимізація VRP',
        '  2.4 Діаграма діяльності (Activity Diagram)',
        '  2.5 Діаграма станів (State Diagram)',
        '  2.6 Діаграма компонентів (Component Diagram)',
        '  2.7 Діаграма розгортання (Deployment Diagram)',
        '3 РОЗРАХУНКОВИЙ РОЗДІЛ',
        '  3.1 Вибір засобів розробки',
        '  3.2 Структура програмного забезпечення',
        '  3.3 Ключові алгоритми',
        'ВИСНОВКИ',
        'ПЕРЕЛІК ДЖЕРЕЛ ПОСИЛАНЬ',
    ]
    for item in items:
        p = doc.add_paragraph()
        fmt_paragraph(p, first_indent=Cm(0))
        run = p.add_run(item)
        fmt_run(run)


def add_vstup(doc):
    """Add ВСТУП (Introduction)."""
    add_page_break(doc)
    add_heading_centered(doc, 'ВСТУП')

    paras = [
        'Сучасна логістика стикається з постійно зростаючим обсягом доставок, '
        'що вимагає ефективних інструментів планування маршрутів. Задача маршрутизації '
        'транспортних засобів (Vehicle Routing Problem, VRP) є однією з ключових задач '
        'комбінаторної оптимізації, що належить до класу NP-складних задач.',

        'Метою даної курсової роботи є проєктування інформаційної системи «DeliveryIQ» '
        'для оптимізації розподілу вантажів у транспортній мережі з використанням '
        'об\'єктно-орієнтованого підходу та мови моделювання UML. Система використовує '
        'реальні дані вуличної мережі OpenStreetMap та підтримує мультимодальну '
        'маршрутизацію (автомобіль, велосипед, пішки).',

        'Об\'єктом проєктування є інформаційна система оптимізації маршрутів '
        'доставки вантажів у міській транспортній мережі.',

        'Предметом проєктування є архітектура, структура класів, взаємодія '
        'компонентів та поведінка системи DeliveryIQ.',

        'Підхід до проєктування: об\'єктно-орієнтований (ООП) з використанням '
        'уніфікованої мови моделювання UML.',
    ]
    for text in paras:
        add_body_paragraph(doc, text)


def add_section1(doc):
    """Section 1: ЗАГАЛЬНИЙ РОЗДІЛ — Аналіз предметної області."""
    add_page_break(doc)
    add_heading_centered(doc, '1 ЗАГАЛЬНИЙ РОЗДІЛ')

    # 1.1
    add_subheading(doc, '1.1 Опис предметної області')

    add_body_paragraph(
        doc,
        'Предметна область \u2014 оптимізація маршрутів доставки вантажів у міській '
        'транспортній мережі. Транспортну мережу формалізовано як мультимодальний '
        'зважений орієнтований граф G = (V, E, \u03a6, \u03a8), де:'
    )

    items = [
        'V = {v\u2081, v\u2082, \u2026, v\u2099} \u2014 скінченна множина вузлів мережі (перехрестя, точки інтересу);',
        'E \u2286 V \u00d7 V \u2014 множина орієнтованих ребер (дорожні сегменти);',
        '\u03a6 : V \u2192 A \u2014 функція атрибутів вузлів (координати, часові вікна);',
        '\u03a8 : E \u2192 B \u2014 функція атрибутів ребер (довжина, допустимі типи транспорту).',
    ]
    for item in items:
        add_bullet_item(doc, item)

    add_body_paragraph(
        doc,
        'Система підтримує три типи транспортних засобів: T = {Car, Bike, Walk}, '
        'кожен з яких має характеристичну швидкість та вантажну місткість. '
        'Кожне ребро графа може бути доступним для одного або кількох типів '
        'транспорту відповідно до тегів OpenStreetMap (highway, motor_vehicle, '
        'bicycle, foot, oneway тощо).'
    )

    add_body_paragraph(
        doc,
        'Задача оптимізації формулюється як задача змішаного цілочислового '
        'програмування (CVRP \u2014 Capacitated Vehicle Routing Problem), де '
        'цільова функція мінімізує зважену суму загального часу доставки '
        'та кількості задіяних транспортних засобів.'
    )

    # 1.2
    add_subheading(doc, '1.2 Огляд існуючих рішень')

    add_body_paragraph(doc, 'На ринку існує ряд рішень для оптимізації маршрутів доставки:')

    solutions = [
        ('Google OR-Tools', 'потужний фреймворк від Google для комбінаторної оптимізації, '
         'але вимагає значного досвіду програмування та не має інтегрованого веб-інтерфейсу;'),
        ('Route4Me, OptimoRoute', 'комерційні SaaS-рішення з широким функціоналом, '
         'але з високою вартістю підписки та відсутністю контролю над алгоритмами;'),
        ('OSRM, Valhalla', 'відкриті маршрутизатори на базі OSM, але без вбудованої '
         'підтримки VRP-оптимізації та мультимодальності в одному інтерфейсі.'),
    ]
    for name, desc in solutions:
        add_bullet_item(doc, desc, bold_prefix=name)

    add_body_paragraph(
        doc,
        'DeliveryIQ відрізняється від існуючих рішень тим, що поєднує: '
        '(1) роботу з реальними даними OSM; (2) мультимодальну маршрутизацію '
        '(drive/bike/walk) в єдиному інтерфейсі; (3) відкритий код та '
        'можливість розгортання на власному сервері; (4) гнучкий вибір '
        'алгоритмів оптимізації (NN, 2-opt, генетичний алгоритм, Christofides).'
    )

    # 1.3
    add_subheading(doc, '1.3 Постановка задачі')

    add_body_paragraph(doc, 'Необхідно спроєктувати інформаційну систему, яка забезпечує:')

    tasks = [
        'геокодування адрес доставки у географічні координати;',
        'завантаження та обробку реальної вуличної мережі з OpenStreetMap;',
        'побудову мультимодальних графів з урахуванням правил доступу для різних типів транспорту;',
        'обчислення матриці відстаней (часів переїзду) за алгоритмом Дейкстри;',
        'розв\'язання задачі комівояжера (TSP) та задачі маршрутизації з обмеженнями на місткість (CVRP);',
        'підтримку механізму «останньої милі» для автомобільного режиму;',
        'візуалізацію оптимальних маршрутів на інтерактивній карті;',
        'керування базою даних пакунків зі статусами доставки.',
    ]
    for i, task in enumerate(tasks, 1):
        add_numbered_item(doc, i, task)


def add_section2(doc):
    """Section 2: СПЕЦІАЛЬНИЙ РОЗДІЛ — Проєктування ІС."""
    add_page_break(doc)
    add_heading_centered(doc, '2 СПЕЦІАЛЬНИЙ РОЗДІЛ')

    add_body_paragraph(
        doc,
        'Для проєктування інформаційної системи DeliveryIQ обрано '
        'об\'єктно-орієнтований підхід з використанням уніфікованої мови '
        'моделювання UML (Unified Modeling Language). Нижче наведено комплект '
        'UML-діаграм, що описують статичну структуру, динамічну поведінку '
        'та фізичну архітектуру системи.'
    )

    # 2.1 Use Case
    add_subheading(doc, '2.1 Діаграма прецедентів (Use Case Diagram)')

    add_body_paragraph(
        doc,
        'Діаграма прецедентів визначає функціональні вимоги системи з точки '
        'зору її користувачів (акторів). Основним актором є Оператор логістики, '
        'який взаємодіє із системою через веб-інтерфейс Streamlit.'
    )

    add_body_paragraph(doc, 'Зовнішні актори-системи:')
    for a in [
        'Nominatim API \u2014 сервіс геокодування OpenStreetMap для перетворення адрес у координати;',
        'OpenStreetMap (Overpass API) \u2014 джерело даних вуличної мережі;',
        'Mapbox API \u2014 опціональний сервіс для отримання матриці часу з урахуванням трафіку.',
    ]:
        add_bullet_item(doc, a)

    add_body_paragraph(doc, 'Основні прецеденти включають:')
    for uc in [
        'обрання міста доставки та задання адреси депо;',
        'додавання точок доставки (текстова адреса або клік на карті);',
        'керування парком транспортних засобів (тип, місткість);',
        'запуск оптимізації маршруту (включає завантаження мережі, побудову матриці, розв\'язання TSP/VRP);',
        'перегляд результатів на інтерактивній карті та у табличному вигляді;',
        'керування базою пакунків.',
    ]:
        add_bullet_item(doc, uc)

    add_image_with_caption(doc, 'UseCase_DeliveryIQ.png',
                           'Рисунок 2.1 \u2013 Діаграма прецедентів системи DeliveryIQ')

    # 2.2 Class diagrams
    add_page_break(doc)
    add_subheading(doc, '2.2 Діаграми класів (Class Diagram)')

    add_body_paragraph(
        doc,
        'Для повного опису структури системи побудовано дві діаграми класів: '
        '(1) діаграма предметної області, що відображає ключові сутності '
        'та їх взаємозв\'язки; (2) діаграма модульної структури, що показує '
        'розподіл функціональності між програмними модулями.'
    )

    # 2.2.1
    add_subheading(doc, '2.2.1 Діаграма класів предметної області')

    add_body_paragraph(doc, 'Основні сутності предметної області:')

    entities = [
        ('Location', 'геокодована адреса з координатами та прив\'язкою до вузла OSM-графа;'),
        ('DeliveryStop', 'точка доставки з адресою, координатами, вагою пакунку та способом додавання;'),
        ('Vehicle', 'транспортний засіб з типом (drive/bike/walk), місткістю та кольором на карті;'),
        ('VehicleRoute', 'маршрут конкретного ТЗ: послідовність зупинок, повний шлях, загальний час;'),
        ('Package', 'фізичний пакунок з адресою, вагою, статусом (PENDING/IN_TRANSIT/DELIVERED);'),
        ('PackageDB', 'SQLite-сховище пакунків з CRUD-операціями;'),
        ('UnreachableStop', 'діагностична сутність для фіксації недосяжних пунктів.'),
    ]
    for name, desc in entities:
        p = doc.add_paragraph()
        fmt_paragraph(p, first_indent=Cm(1.25))
        r1 = p.add_run(f'– {name}')
        fmt_run(r1, bold=True, font_name='Consolas', size=Pt(12))
        r2 = p.add_run(f' \u2014 {desc}')
        fmt_run(r2)

    add_image_with_caption(doc, 'Class_Domain_DeliveryIQ.png',
                           'Рисунок 2.2 \u2013 Діаграма класів предметної області')

    # 2.2.2
    add_page_break(doc)
    add_subheading(doc, '2.2.2 Діаграма класів модульної структури')

    add_body_paragraph(
        doc,
        'Система складається з шести основних модулів, кожен з яких реалізує '
        'окрему відповідальність:'
    )

    modules = [
        ('app.py', 'головний модуль Streamlit UI \u2014 координація всіх операцій, керування сесією, кешування;'),
        ('geocoder.py', 'геокодування адрес через Nominatim (forward/reverse);'),
        ('graph_builder.py', 'завантаження OSM-мережі, LSCC-обрізка, побудова модальних графів з travel_time атрибутами;'),
        ('route_solver.py', 'побудова матриці відстаней (Dijkstra), розв\'язання TSP (NN, 2-opt, генетичний алгоритм, Christofides);'),
        ('vrp_solver.py', 'CVRP-оптимізація: розподіл зупинок між ТЗ, K-Means кластеризація, TSP для кожного кластера;'),
        ('visualizer.py', 'побудова інтерактивної карти Folium з AntPath-анімацією.'),
    ]
    for name, desc in modules:
        p = doc.add_paragraph()
        fmt_paragraph(p, first_indent=Cm(1.25))
        r1 = p.add_run(f'– {name}')
        fmt_run(r1, bold=True, font_name='Consolas', size=Pt(12))
        r2 = p.add_run(f' \u2014 {desc}')
        fmt_run(r2)

    add_image_with_caption(doc, 'Class_Modules_DeliveryIQ.png',
                           'Рисунок 2.3 \u2013 Діаграма класів модульної структури')

    # 2.3 Sequence diagrams
    add_page_break(doc)
    add_subheading(doc, '2.3 Діаграми послідовності (Sequence Diagram)')

    add_body_paragraph(
        doc,
        'Діаграми послідовності описують динамічну взаємодію об\'єктів системи '
        'у часі для ключових сценаріїв використання.'
    )

    # 2.3.1
    add_subheading(doc, '2.3.1 Процес оптимізації маршруту (TSP)')

    add_body_paragraph(
        doc,
        'Основний сценарій системи \u2014 повний цикл оптимізації маршруту \u2014 '
        'складається з шести послідовних фаз:'
    )

    phases = [
        'Завантаження OSM-мережі: graph_from_point \u2192 LSCC-обрізка;',
        'Прив\'язка адрес до вузлів графа (nearest_node) з дедуплікацією;',
        'Побудова модальних графів (drive/bike/walk) з travel_time атрибутами;',
        'Побудова матриці відстаней (Dijkstra all-pairs) для кожного режиму;',
        'Розв\'язання TSP: автоматичний вибір методу за кількістю зупинок;',
        'Візуалізація результатів на карті Folium з AntPath-анімацією.',
    ]
    for i, phase in enumerate(phases, 1):
        add_numbered_item(doc, i, phase)

    add_image_with_caption(doc, 'Sequence_Optimization.png',
                           'Рисунок 2.4 \u2013 Діаграма послідовності: оптимізація маршруту TSP')

    # 2.3.2
    add_page_break(doc)
    add_subheading(doc, '2.3.2 Геокодування адреси доставки')

    add_body_paragraph(
        doc,
        'Діаграма описує два альтернативні сценарії додавання точки доставки: '
        '(1) введення текстової адреси з city-lock механізмом та forward-геокодуванням; '
        '(2) клік на карті з reverse-геокодуванням координат.'
    )

    add_image_with_caption(doc, 'Sequence_Geocoding.png',
                           'Рисунок 2.5 \u2013 Діаграма послідовності: геокодування адреси',
                           width=Cm(14))

    # 2.3.3
    add_subheading(doc, '2.3.3 Оптимізація VRP (декілька транспортних засобів)')

    add_body_paragraph(
        doc,
        'Сценарій VRP-оптимізації включає три фази: (1) призначення зупинок '
        'до типів ТЗ за критерієм досяжності та залишкової місткості; '
        '(2) географічна кластеризація K-Means всередині кожного типу; '
        '(3) незалежне розв\'язання TSP для кожного транспортного засобу.'
    )

    add_image_with_caption(doc, 'Sequence_VRP.png',
                           'Рисунок 2.6 \u2013 Діаграма послідовності: VRP-оптимізація',
                           width=Cm(14))

    # 2.4 Activity
    add_page_break(doc)
    add_subheading(doc, '2.4 Діаграма діяльності (Activity Diagram)')

    add_body_paragraph(
        doc,
        'Діаграма діяльності моделює послідовність дій під час виконання '
        'повного циклу оптимізації маршруту. Вона відображає:'
    )

    for item in [
        'послідовні кроки обробки (завантаження графа, прив\'язка, дедуплікація);',
        'умовні розгалуження (перевірка на порожній граф, недостатню кількість вузлів);',
        'паралельні потоки (побудова матриць та розв\'язання TSP для трьох режимів одночасно);',
        'точки синхронізації (fork/join).',
    ]:
        add_bullet_item(doc, item)

    add_image_with_caption(doc, 'Activity_Optimization.png',
                           'Рисунок 2.7 \u2013 Діаграма діяльності: процес оптимізації маршруту',
                           width=Cm(14))

    # 2.5 State
    add_page_break(doc)
    add_subheading(doc, '2.5 Діаграма станів (State Diagram)')

    add_body_paragraph(
        doc,
        'Діаграма станів описує життєвий цикл сутності Package (пакунок) '
        'у базі даних системи. Пакунок може перебувати в одному з трьох станів:'
    )

    states = [
        ('PENDING (Очікує)', 'початковий стан після створення; адреса задана, координати ще не геокодовані;'),
        ('IN_TRANSIT (В дорозі)', 'пакунок включено в оптимізований маршрут ТЗ, координати геокодовані;'),
        ('DELIVERED (Доставлено)', 'пакунок успішно доставлено на адресу, маршрут завершено.'),
    ]
    for name, desc in states:
        add_bullet_item(doc, desc, bold_prefix=name)

    add_body_paragraph(
        doc,
        'Переходи між станами ініціюються операціями set_status() класу PackageDB. '
        'Передбачено зворотний перехід IN_TRANSIT \u2192 PENDING у разі збою доставки.'
    )

    add_image_with_caption(doc, 'State_Package.png',
                           'Рисунок 2.8 \u2013 Діаграма станів: життєвий цикл пакунку',
                           width=Cm(14))

    # 2.6 Component
    add_page_break(doc)
    add_subheading(doc, '2.6 Діаграма компонентів (Component Diagram)')

    add_body_paragraph(
        doc,
        'Діаграма компонентів відображає фізичну структуру системи на рівні '
        'програмних модулів та їх залежностей. Система DeliveryIQ складається з:'
    )

    for item in [
        '6 основних Python-модулів (app.py, geocoder.py, graph_builder.py, route_solver.py, vrp_solver.py, visualizer.py);',
        'модуля package_db.py для роботи з SQLite базою даних packages.db;',
        '6 зовнішніх Python-бібліотек (Streamlit, OSMnx, NetworkX, Folium, geopy, scikit-learn);',
        '3 зовнішніх API-сервісів (Nominatim, Overpass, Mapbox).',
    ]:
        add_bullet_item(doc, item)

    add_image_with_caption(doc, 'Component_DeliveryIQ.png',
                           'Рисунок 2.9 \u2013 Діаграма компонентів системи DeliveryIQ')

    # 2.7 Deployment
    add_page_break(doc)
    add_subheading(doc, '2.7 Діаграма розгортання (Deployment Diagram)')

    add_body_paragraph(doc, 'Діаграма розгортання показує фізичну топологію розгортання системи:')

    for item in [
        'Клієнтський пристрій: веб-браузер (Chrome/Firefox), що взаємодіє зі Streamlit-сервером через HTTP/WebSocket;',
        'Сервер додатку: Python 3.11+ Runtime із Streamlit Server (порт 8501), '
        'усіма модулями системи, SQLite базою даних та кешем GraphML файлів;',
        'Зовнішні хмарні сервіси: OpenStreetMap (Nominatim + Overpass API) та Mapbox.',
    ]:
        add_bullet_item(doc, item)

    add_body_paragraph(
        doc,
        'Кешування OSM-графів у форматі GraphML дозволяє уникнути повторних '
        'звернень до Overpass API та значно пришвидшує повторні запуски оптимізації.'
    )

    add_image_with_caption(doc, 'Deployment_DeliveryIQ.png',
                           'Рисунок 2.10 \u2013 Діаграма розгортання системи DeliveryIQ',
                           width=Cm(14))


def add_section3(doc):
    """Section 3: РОЗРАХУНКОВИЙ РОЗДІЛ — Реалізація."""
    add_page_break(doc)
    add_heading_centered(doc, '3 РОЗРАХУНКОВИЙ РОЗДІЛ')

    # 3.1
    add_subheading(doc, '3.1 Вибір засобів розробки')

    add_body_paragraph(doc, 'Для реалізації системи обрано наступний технологічний стек:')

    tech = [
        ('Python 3.11+', 'основна мова програмування;'),
        ('Streamlit', 'фреймворк для побудови інтерактивного веб-інтерфейсу;'),
        ('OSMnx 2.x', 'бібліотека для завантаження та аналізу мереж OpenStreetMap;'),
        ('NetworkX', 'бібліотека для роботи з графами (Dijkstra, побудова маршрутів);'),
        ('Folium + AntPath', 'інтерактивна картографічна візуалізація з анімованими маршрутами;'),
        ('geopy (Nominatim)', 'геокодування адрес у координати;'),
        ('scikit-learn (KMeans)', 'географічна кластеризація для VRP;'),
        ('SQLite', 'легка реляційна БД для зберігання пакунків;'),
        ('PlantUML', 'засіб побудови UML-діаграм для проєктування.'),
    ]
    for name, desc in tech:
        add_bullet_item(doc, desc, bold_prefix=name)

    # 3.2
    add_subheading(doc, '3.2 Структура програмного забезпечення')

    add_body_paragraph(
        doc,
        'Архітектура системи побудована за принципом чіткого розділення '
        'відповідальностей (Separation of Concerns). Конвеєр обробки даних:'
    )

    pipeline = [
        'app.py (Streamlit UI) \u2014 точка входу, координація всіх операцій;',
        '\u2192 geocoder.py \u2014 перетворення адрес у координати (Nominatim);',
        '\u2192 graph_builder.py \u2014 завантаження OSM-мережі, LSCC-обрізка, модальні графи;',
        '\u2192 route_solver.py \u2014 матриця відстаней (Dijkstra) + TSP (NN / 2-opt / GA / Christofides);',
        '\u2192 vrp_solver.py \u2014 CVRP: розподіл зупинок, K-Means, TSP для кожного ТЗ;',
        '\u2192 visualizer.py \u2014 Folium-карта з AntPath-маршрутами;',
        '\u2192 package_db.py \u2014 SQLite CRUD для пакунків.',
    ]
    for item in pipeline:
        add_body_paragraph(doc, item, first_indent=Cm(1.25))

    # 3.3
    add_subheading(doc, '3.3 Ключові алгоритми')

    add_body_paragraph(doc, 'Система реалізує декілька алгоритмів оптимізації маршрутів:')

    algorithms = [
        ('Nearest Neighbour (NN)', 'жадібний алгоритм O(n\u00b2), використовується для 1\u20132 зупинок;'),
        ('2-opt', 'ітеративне покращення NN-розв\'язку шляхом обміну ребер, O(n\u00b3) за ітерацію;'),
        ('Генетичний алгоритм (GA)', 'Order Crossover (OX) з елітизмом, 120 особин \u00d7 400 поколінь, для задач із >20 зупинками;'),
        ('Christofides', '\u00bd-наближення для метричного TSP, потребує повний зв\'язний граф.'),
    ]
    for name, desc in algorithms:
        add_bullet_item(doc, desc, bold_prefix=name)

    add_body_paragraph(
        doc,
        'Автоматичний вибір алгоритму (method="auto") визначається кількістю '
        'зупинок: n\u22642 \u2192 NN, 3\u2264n\u226420 \u2192 2-opt, n>20 \u2192 генетичний алгоритм. '
        'Це забезпечує баланс між якістю розв\'язку та часом обчислення.'
    )

    add_body_paragraph(
        doc,
        'Механізм «останньої милі» для автомобільного режиму: якщо адреса '
        'доставки знаходиться на пішохідній вулиці, автомобіль паркується '
        'на найближчому car-accessible вузлі, а залишок шляху (до 100 м) '
        'долається пішки. Час переїзду включає обидві частини.'
    )


def add_vysnovky(doc):
    """ВИСНОВКИ (Conclusions)."""
    add_page_break(doc)
    add_heading_centered(doc, 'ВИСНОВКИ')

    add_body_paragraph(
        doc,
        'В результаті курсової роботи було спроєктовано інформаційну '
        'систему «DeliveryIQ» для оптимізації розподілу вантажів у транспортній '
        'мережі з використанням об\'єктно-орієнтованого підходу.'
    )

    add_body_paragraph(doc, 'В ході роботи було:')

    conclusions = [
        'проведено аналіз предметної області маршрутизації транспортних засобів та формалізовано задачу CVRP;',
        'побудовано діаграму прецедентів, що визначає функціональні вимоги системи та її акторів;',
        'розроблено дві діаграми класів: предметної області (ключові сутності та зв\'язки) та модульної структури (розподіл функціональності між модулями);',
        'створено три діаграми послідовності для основних сценаріїв: оптимізація TSP, геокодування адрес, VRP-оптимізація;',
        'побудовано діаграму діяльності процесу оптимізації з паралельними потоками;',
        'розроблено діаграму станів для життєвого циклу пакунку;',
        'побудовано діаграми компонентів та розгортання, що описують фізичну архітектуру системи;',
        'описано технологічний стек та ключові алгоритми оптимізації.',
    ]
    for i, item in enumerate(conclusions, 1):
        add_numbered_item(doc, i, item)

    add_body_paragraph(
        doc,
        'Загалом побудовано 10 UML-діаграм, що забезпечують повне покриття '
        'статичної структури, динамічної поведінки та фізичної архітектури '
        'системи DeliveryIQ. Розроблений комплект проєктної документації '
        'може бути використаний як основа для подальшої реалізації, '
        'тестування та супроводу системи.'
    )

    add_body_paragraph(
        doc,
        'Пояснювальну записку до курсової роботи оформлено у відповідності до '
        'методичних рекомендацій до змісту та структури курсової роботи з дисципліни '
        '«Проєктування інформаційних систем» за освітнім рівнем «бакалавр» для '
        'студентів спеціальності 122 Комп\'ютерні науки [10].'
    )


def add_references(doc):
    """ПЕРЕЛІК ДЖЕРЕЛ ПОСИЛАНЬ."""
    add_page_break(doc)
    add_heading_centered(doc, 'ПЕРЕЛІК ДЖЕРЕЛ ПОСИЛАНЬ')

    sources = [
        'Шаховська Н. Б. Проектування інформаційних систем: Навчальний посібник / Н. Б. Шаховська, В. В. Литвин; за ред. В. В. Пасічника. \u2013 Львів: «Магнолія 2006», 2017. \u2013 380 с.',
        'Проектування інформаційних систем: Посібник / За ред. В. С. Пономаренка. \u2013 К.: Академія, 2002. \u2013 488 с.',
        'Гайна Г. А. Основи проектування баз даних: Навчальний посібник. \u2013 К.: Кондор, 2008. \u2013 200 с.',
        'Зеленцов Д.Г., Ляшенко О.А., Науменко Н.Ю. Информационное обеспечение расчетов корродирующих объектов. Математические модели и концепция проектирования систем. (монографія) // Днепропетровск: УГХТУ, 2012. \u2013 264 с. DOI: 10.32434/mono-1-ZDG-LOA-NNY',
        'Авраменко В. С. Проектування інформаційних систем: навчальний посібник / В. С. Авраменко, А. С. Авраменко. \u2013 Черкаси: Черкаський національний університет ім. Б. Хмельницького, 2017. \u2013 434 с.',
        'Boeing G. OSMnx: New methods for acquiring, constructing, analyzing, and visualizing complex street networks. Computers, Environment and Urban Systems, 2017. Vol. 65. P. 126\u2013139.',
        'Toth P., Vigo D. The Vehicle Routing Problem. SIAM Monographs on Discrete Mathematics and Applications, 2002.',
        'Christofides N. Worst-case analysis of a new heuristic for the travelling salesman problem. Report 388, Graduate School of Industrial Administration, Carnegie Mellon University, 1976.',
        'Goldberg D.E. Genetic Algorithms in Search, Optimization, and Machine Learning. Addison-Wesley, 1989.',
        'Методичні вказівки до виконання курсових робіт з дисципліни «Проектування інформаційних систем» за освітнім рівнем «бакалавр» для студентів спеціальності 122 «Комп\'ютерні науки» / Укл. О.А. Ляшенко. \u2013 Дніпро: ДВНЗ «УДХТУ», 2020. \u2013 56 с.',
    ]
    for i, src in enumerate(sources, 1):
        p = doc.add_paragraph()
        fmt_paragraph(p, first_indent=Cm(1.25))
        run = p.add_run(f'{i}. {src}')
        fmt_run(run)


# ── Main ────────────────────────────────────────────────────────────────────

def main():
    print(f"Loading template: {TEMPLATE}")
    doc = Document(str(TEMPLATE))

    print("Updating assignment table...")
    update_assignment_table(doc)

    print("Stripping placeholder body content...")
    strip_body_after_tables(doc)

    print("Adding Реферат...")
    add_referat(doc, skip_page_break=True)  # sectPr from template already starts new page

    print("Adding ЗМІСТ...")
    add_toc(doc)

    print("Adding ВСТУП...")
    add_vstup(doc)

    print("Adding 1 ЗАГАЛЬНИЙ РОЗДІЛ...")
    add_section1(doc)

    print("Adding 2 СПЕЦІАЛЬНИЙ РОЗДІЛ...")
    add_section2(doc)

    print("Adding 3 РОЗРАХУНКОВИЙ РОЗДІЛ...")
    add_section3(doc)

    print("Adding ВИСНОВКИ...")
    add_vysnovky(doc)

    print("Adding ПЕРЕЛІК ДЖЕРЕЛ ПОСИЛАНЬ...")
    add_references(doc)

    print(f"Saving to: {OUTPUT}")
    doc.save(str(OUTPUT))
    print("Done!")


if __name__ == '__main__':
    main()
