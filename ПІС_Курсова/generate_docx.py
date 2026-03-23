"""
Generate the coursework document for ПІС (Проектування Інформаційних Систем).
Topic: DeliveryIQ — Система оптимізації маршрутів доставки вантажів.
Approach: OOP (UML diagrams).
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

BASE = Path(__file__).parent
DIAGRAMS = BASE / "diagrams"
OUTPUT = BASE / "Курсова_робота_ПІС.docx"


def set_style(doc):
    """Configure base styles for the document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    for level in range(1, 4):
        sname = f'Heading {level}'
        s = doc.styles[sname]
        s.font.name = 'Times New Roman'
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.font.bold = True
        if level == 1:
            s.font.size = Pt(16)
        elif level == 2:
            s.font.size = Pt(14)
        else:
            s.font.size = Pt(14)
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.5


def add_title_page(doc):
    doc.add_paragraph('')  # spacer
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ')
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('НАЦІОНАЛЬНИЙ ТЕХНІЧНИЙ УНІВЕРСИТЕТ')
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('КУРСОВА РОБОТА')
    run.font.size = Pt(18)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('з дисципліни «Проектування інформаційних систем»')
    run.font.size = Pt(14)

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('на тему:')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('«Проектування інформаційної системи оптимізації\n'
                     'розподілу вантажів у транспортній мережі (DeliveryIQ)»')
    run.font.size = Pt(16)
    run.bold = True

    for _ in range(6):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Виконав: студент групи ___\n'
                     '_________________________\n\n'
                     'Перевірив: ________________\n'
                     '_________________________')
    run.font.size = Pt(14)

    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2025')
    run.font.size = Pt(14)

    doc.add_page_break()


def add_toc(doc):
    doc.add_heading('ЗМІСТ', level=1)
    items = [
        ('ВСТУП', 3),
        ('1. АНАЛІЗ ПРЕДМЕТНОЇ ОБЛАСТІ', 4),
        ('1.1. Опис предметної області', 4),
        ('1.2. Огляд існуючих рішень', 5),
        ('1.3. Постановка задачі', 5),
        ('2. ПРОЕКТУВАННЯ ІНФОРМАЦІЙНОЇ СИСТЕМИ', 6),
        ('2.1. Діаграма прецедентів (Use Case)', 6),
        ('2.2. Діаграми класів (Class Diagram)', 7),
        ('2.3. Діаграми послідовності (Sequence Diagram)', 8),
        ('2.4. Діаграма діяльності (Activity Diagram)', 9),
        ('2.5. Діаграма станів (State Diagram)', 10),
        ('2.6. Діаграма компонентів (Component Diagram)', 11),
        ('2.7. Діаграма розгортання (Deployment Diagram)', 12),
        ('3. РЕАЛІЗАЦІЯ', 13),
        ('3.1. Вибір засобів розробки', 13),
        ('3.2. Структура програмного забезпечення', 13),
        ('3.3. Ключові алгоритми', 14),
        ('ВИСНОВКИ', 15),
        ('СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ', 16),
    ]
    for title, page in items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}')
        run.font.size = Pt(14)
    doc.add_page_break()


def add_image(doc, name, caption, width=Cm(16)):
    """Add a diagram image with caption."""
    img_path = DIAGRAMS / name
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=width)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.size = Pt(12)
    run.italic = True
    doc.add_paragraph('')  # spacer


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    set_style(doc)
    add_title_page(doc)
    add_toc(doc)

    # =========================================================================
    # ВСТУП
    # =========================================================================
    doc.add_heading('ВСТУП', level=1)

    doc.add_paragraph(
        'Сучасна логістика стикається з постійно зростаючим обсягом доставок, '
        'що вимагає ефективних інструментів планування маршрутів. Задача маршрутизації '
        'транспортних засобів (Vehicle Routing Problem, VRP) є однією з ключових задач '
        'комбінаторної оптимізації, що належить до класу NP-складних задач.'
    )

    doc.add_paragraph(
        'Метою даної курсової роботи є проектування інформаційної системи «DeliveryIQ» '
        'для оптимізації розподілу вантажів у транспортній мережі з використанням '
        'об\'єктно-орієнтованого підходу та мови моделювання UML. Система використовує '
        'реальні дані вуличної мережі OpenStreetMap та підтримує мультимодальну '
        'маршрутизацію (автомобіль, велосипед, пішки).'
    )

    doc.add_paragraph(
        'Об\'єктом проектування є інформаційна система оптимізації маршрутів '
        'доставки вантажів у міській транспортній мережі.'
    )

    doc.add_paragraph(
        'Предметом проектування є архітектура, структура класів, взаємодія '
        'компонентів та поведінка системи DeliveryIQ.'
    )

    doc.add_paragraph(
        'Підхід до проектування: об\'єктно-орієнтований (ООП) з використанням '
        'уніфікованої мови моделювання UML.'
    )

    doc.add_page_break()

    # =========================================================================
    # РОЗДІЛ 1
    # =========================================================================
    doc.add_heading('1. АНАЛІЗ ПРЕДМЕТНОЇ ОБЛАСТІ', level=1)

    doc.add_heading('1.1. Опис предметної області', level=2)

    doc.add_paragraph(
        'Предметна область — оптимізація маршрутів доставки вантажів у міській '
        'транспортній мережі. Транспортну мережу формалізовано як мультимодальний '
        'зважений орієнтований граф G = (V, E, Φ, Ψ), де:'
    )

    items = [
        'V = {v₁, v₂, …, vₙ} — скінченна множина вузлів мережі (перехрестя, точки інтересу);',
        'E ⊆ V × V — множина орієнтованих ребер (дорожні сегменти);',
        'Φ : V → A — функція атрибутів вузлів (координати, часові вікна);',
        'Ψ : E → B — функція атрибутів ребер (довжина, допустимі типи транспорту).',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'Система підтримує три типи транспортних засобів: T = {Car, Bike, Walk}, '
        'кожен з яких має характеристичну швидкість та вантажну місткість. '
        'Кожне ребро графа може бути доступним для одного або кількох типів '
        'транспорту відповідно до тегів OpenStreetMap (highway, motor_vehicle, '
        'bicycle, foot, oneway тощо).'
    )

    doc.add_paragraph(
        'Задача оптимізації формулюється як задача змішаного цілочислового '
        'програмування (CVRP — Capacitated Vehicle Routing Problem), де '
        'цільова функція мінімізує зважену суму загального часу доставки '
        'та кількості задіяних транспортних засобів.'
    )

    doc.add_heading('1.2. Огляд існуючих рішень', level=2)

    doc.add_paragraph(
        'На ринку існує ряд рішень для оптимізації маршрутів доставки:'
    )

    solutions = [
        ('Google OR-Tools', 'потужний фреймворк від Google для комбінаторної оптимізації, '
         'але вимагає значного досвіду програмування та не має інтегрованого веб-інтерфейсу;'),
        ('Route4Me, OptimoRoute', 'комерційні SaaS-рішення з широким функціоналом, '
         'але з високою вартістю підписки та відсутністю контролю над алгоритмами;'),
        ('OSRM, Valhalla', 'відкриті маршрутизатори на базі OSM, але без вбудованої '
         'підтримки VRP-оптимізації та мультимодальності в одному інтерфейсі.'),
    ]
    for name, desc in solutions:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name)
        run.bold = True
        p.add_run(f' — {desc}')

    doc.add_paragraph(
        'DeliveryIQ відрізняється від існуючих рішень тим, що поєднує: '
        '(1) роботу з реальними даними OSM; (2) мультимодальну маршрутизацію '
        '(drive/bike/walk) в єдиному інтерфейсі; (3) відкритий код та '
        'можливість розгортання на власному сервері; (4) гнучкий вибір '
        'алгоритмів оптимізації (NN, 2-opt, генетичний алгоритм, Christofides).'
    )

    doc.add_heading('1.3. Постановка задачі', level=2)

    doc.add_paragraph('Необхідно спроектувати інформаційну систему, яка забезпечує:')

    tasks = [
        'геокодування адрес доставки у географічні координати;',
        'завантаження та обробку реальної вуличної мережі з OpenStreetMap;',
        'побудову мультимодальних графів з урахуванням правил доступу для різних типів транспорту;',
        'обчислення матриці відстаней (часів переїзду) за алгоритмом Дейкстри;',
        'розв\'язання задачі комівояжера (TSP) та задачі маршрутизації з обмеженнями на місткість (CVRP);',
        'підтримку механізму «останньої милі» для автомобільного режиму;',
        'візуалізацію оптимальних маршрутів на інтерактивній карті;',
        'керування базою даних пакунків зі статусами доставки.',
    ]
    for i, task in enumerate(tasks, 1):
        doc.add_paragraph(f'{i}) {task}')

    doc.add_page_break()

    # =========================================================================
    # РОЗДІЛ 2
    # =========================================================================
    doc.add_heading('2. ПРОЕКТУВАННЯ ІНФОРМАЦІЙНОЇ СИСТЕМИ', level=1)

    doc.add_paragraph(
        'Для проектування інформаційної системи DeliveryIQ обрано '
        'об\'єктно-орієнтований підхід з використанням уніфікованої мови '
        'моделювання UML (Unified Modeling Language). Нижче наведено комплект '
        'UML-діаграм, що описують статичну структуру, динамічну поведінку '
        'та фізичну архітектуру системи.'
    )

    # 2.1 Use Case
    doc.add_heading('2.1. Діаграма прецедентів (Use Case Diagram)', level=2)

    doc.add_paragraph(
        'Діаграма прецедентів визначає функціональні вимоги системи з точки '
        'зору її користувачів (акторів). Основним актором є Оператор логістики, '
        'який взаємодіє із системою через веб-інтерфейс Streamlit.'
    )

    doc.add_paragraph('Зовнішні актори-системи:')
    ext_actors = [
        'Nominatim API — сервіс геокодування OpenStreetMap для перетворення адрес у координати;',
        'OpenStreetMap (Overpass API) — джерело даних вуличної мережі;',
        'Mapbox API — опціональний сервіс для отримання матриці часу з урахуванням трафіку.',
    ]
    for a in ext_actors:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_paragraph('Основні прецеденти включають:')
    usecases = [
        'обрання міста доставки та задання адреси депо;',
        'додавання точок доставки (текстова адреса або клік на карті);',
        'керування парком транспортних засобів (тип, місткість);',
        'запуск оптимізації маршруту (включає завантаження мережі, побудову матриці, розв\'язання TSP/VRP);',
        'перегляд результатів на інтерактивній карті та у табличному вигляді;',
        'керування базою пакунків.',
    ]
    for uc in usecases:
        doc.add_paragraph(uc, style='List Bullet')

    add_image(doc, 'UseCase_DeliveryIQ.png',
              'Рис. 2.1 — Діаграма прецедентів системи DeliveryIQ')

    doc.add_page_break()

    # 2.2 Class diagrams
    doc.add_heading('2.2. Діаграми класів (Class Diagram)', level=2)

    doc.add_paragraph(
        'Для повного опису структури системи побудовано дві діаграми класів: '
        '(1) діаграма предметної області, що відображає ключові сутності '
        'та їх взаємозв\'язки; (2) діаграма модульної структури, що показує '
        'розподіл функціональності між програмними модулями.'
    )

    doc.add_heading('2.2.1. Діаграма класів предметної області', level=3)

    doc.add_paragraph(
        'Основні сутності предметної області:'
    )

    entities = [
        ('Location', 'геокодована адреса з координатами та прив\'язкою до вузла OSM-графа;'),
        ('DeliveryStop', 'точка доставки з адресою, координатами, вагою пакунку та способом додавання;'),
        ('Vehicle', 'транспортний засіб з типом (drive/bike/walk), місткістю та кольором на карті;'),
        ('VehicleRoute', 'маршрут конкретного ТЗ: послідовність зупинок, повний шлях, загальний час;'),
        ('Package', 'фізичний пакунок з адресою, вагою, статусом (PENDING/IN_TRANSIT/DELIVERED);'),
        ('PackageDB', 'SQLite-сховище пакунків з CRUD-операціями;'),
        ('UnreachableStop', 'діагностична сутність для фіксації недосяжних пунктів.'),
    ]
    for name, desc in entities:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name)
        run.bold = True
        run.font.name = 'Consolas'
        p.add_run(f' — {desc}')

    add_image(doc, 'Class_Domain_DeliveryIQ.png',
              'Рис. 2.2 — Діаграма класів предметної області', width=Cm(16))

    doc.add_page_break()

    doc.add_heading('2.2.2. Діаграма класів модульної структури', level=3)

    doc.add_paragraph(
        'Система складається з шести основних модулів, кожен з яких реалізує '
        'окрему відповідальність:'
    )

    modules = [
        ('app.py', 'головний модуль Streamlit UI — координація всіх операцій, '
         'керування сесією, кешування;'),
        ('geocoder.py', 'геокодування адрес через Nominatim (forward/reverse);'),
        ('graph_builder.py', 'завантаження OSM-мережі, LSCC-обрізка, побудова '
         'модальних графів з travel_time атрибутами;'),
        ('route_solver.py', 'побудова матриці відстаней (Dijkstra), розв\'язання '
         'TSP (NN, 2-opt, генетичний алгоритм, Christofides);'),
        ('vrp_solver.py', 'CVRP-оптимізація: розподіл зупинок між ТЗ, '
         'K-Means кластеризація, TSP для кожного кластера;'),
        ('visualizer.py', 'побудова інтерактивної карти Folium з AntPath-анімацією.'),
    ]
    for name, desc in modules:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name)
        run.bold = True
        run.font.name = 'Consolas'
        p.add_run(f' — {desc}')

    add_image(doc, 'Class_Modules_DeliveryIQ.png',
              'Рис. 2.3 — Діаграма класів модульної структури', width=Cm(16))

    doc.add_page_break()

    # 2.3 Sequence diagrams
    doc.add_heading('2.3. Діаграми послідовності (Sequence Diagram)', level=2)

    doc.add_paragraph(
        'Діаграми послідовності описують динамічну взаємодію об\'єктів системи '
        'у часі для ключових сценаріїв використання.'
    )

    doc.add_heading('2.3.1. Процес оптимізації маршруту (TSP)', level=3)

    doc.add_paragraph(
        'Основний сценарій системи — повний цикл оптимізації маршруту — '
        'складається з шести послідовних фаз:'
    )

    phases = [
        'Завантаження OSM-мережі: graph_from_point → LSCC-обрізка;',
        'Прив\'язка адрес до вузлів графа (nearest_node) з дедуплікацією;',
        'Побудова модальних графів (drive/bike/walk) з travel_time атрибутами;',
        'Побудова матриці відстаней (Dijkstra all-pairs) для кожного режиму;',
        'Розв\'язання TSP: автоматичний вибір методу за кількістю зупинок;',
        'Візуалізація результатів на карті Folium з AntPath-анімацією.',
    ]
    for i, phase in enumerate(phases, 1):
        doc.add_paragraph(f'{i}. {phase}')

    add_image(doc, 'Sequence_Optimization.png',
              'Рис. 2.4 — Діаграма послідовності: оптимізація маршруту TSP', width=Cm(16))

    doc.add_page_break()

    doc.add_heading('2.3.2. Геокодування адреси доставки', level=3)

    doc.add_paragraph(
        'Діаграма описує два альтернативні сценарії додавання точки доставки: '
        '(1) введення текстової адреси з city-lock механізмом та forward-геокодуванням; '
        '(2) клік на карті з reverse-геокодуванням координат.'
    )

    add_image(doc, 'Sequence_Geocoding.png',
              'Рис. 2.5 — Діаграма послідовності: геокодування адреси', width=Cm(14))

    doc.add_heading('2.3.3. Оптимізація VRP (декілька транспортних засобів)', level=3)

    doc.add_paragraph(
        'Сценарій VRP-оптимізації включає три фази: (1) призначення зупинок '
        'до типів ТЗ за критерієм досяжності та залишкової місткості; '
        '(2) географічна кластеризація K-Means всередині кожного типу; '
        '(3) незалежне розв\'язання TSP для кожного транспортного засобу.'
    )

    add_image(doc, 'Sequence_VRP.png',
              'Рис. 2.6 — Діаграма послідовності: VRP-оптимізація', width=Cm(14))

    doc.add_page_break()

    # 2.4 Activity
    doc.add_heading('2.4. Діаграма діяльності (Activity Diagram)', level=2)

    doc.add_paragraph(
        'Діаграма діяльності моделює послідовність дій під час виконання '
        'повного циклу оптимізації маршруту. Вона відображає:'
    )

    activity_items = [
        'послідовні кроки обробки (завантаження графа, прив\'язка, дедуплікація);',
        'умовні розгалуження (перевірка на порожній граф, недостатню кількість вузлів);',
        'паралельні потоки (побудова матриць та розв\'язання TSP для трьох режимів одночасно);',
        'точки синхронізації (fork/join).',
    ]
    for item in activity_items:
        doc.add_paragraph(item, style='List Bullet')

    add_image(doc, 'Activity_Optimization.png',
              'Рис. 2.7 — Діаграма діяльності: процес оптимізації маршруту', width=Cm(14))

    doc.add_page_break()

    # 2.5 State
    doc.add_heading('2.5. Діаграма станів (State Diagram)', level=2)

    doc.add_paragraph(
        'Діаграма станів описує життєвий цикл сутності Package (пакунок) '
        'у базі даних системи. Пакунок може перебувати в одному з трьох станів:'
    )

    states = [
        ('PENDING (Очікує)', 'початковий стан після створення; адреса задана, координати ще не геокодовані;'),
        ('IN_TRANSIT (В дорозі)', 'пакунок включено в оптимізований маршрут ТЗ, координати геокодовані;'),
        ('DELIVERED (Доставлено)', 'пакунок успішно доставлено на адресу, маршрут завершено.'),
    ]
    for name, desc in states:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name)
        run.bold = True
        p.add_run(f' — {desc}')

    doc.add_paragraph(
        'Переходи між станами ініціюються операціями set_status() класу PackageDB. '
        'Передбачено зворотний перехід IN_TRANSIT → PENDING у разі збою доставки.'
    )

    add_image(doc, 'State_Package.png',
              'Рис. 2.8 — Діаграма станів: життєвий цикл пакунку', width=Cm(14))

    doc.add_page_break()

    # 2.6 Component
    doc.add_heading('2.6. Діаграма компонентів (Component Diagram)', level=2)

    doc.add_paragraph(
        'Діаграма компонентів відображає фізичну структуру системи на рівні '
        'програмних модулів та їх залежностей. Система DeliveryIQ складається з:'
    )

    comp_items = [
        '6 основних Python-модулів (app.py, geocoder.py, graph_builder.py, route_solver.py, vrp_solver.py, visualizer.py);',
        'модуля package_db.py для роботи з SQLite базою даних packages.db;',
        '6 зовнішніх Python-бібліотек (Streamlit, OSMnx, NetworkX, Folium, geopy, scikit-learn);',
        '3 зовнішніх API-сервісів (Nominatim, Overpass, Mapbox).',
    ]
    for item in comp_items:
        doc.add_paragraph(item, style='List Bullet')

    add_image(doc, 'Component_DeliveryIQ.png',
              'Рис. 2.9 — Діаграма компонентів системи DeliveryIQ', width=Cm(16))

    doc.add_page_break()

    # 2.7 Deployment
    doc.add_heading('2.7. Діаграма розгортання (Deployment Diagram)', level=2)

    doc.add_paragraph(
        'Діаграма розгортання показує фізичну топологію розгортання системи:'
    )

    deploy_items = [
        'Клієнтський пристрій: веб-браузер (Chrome/Firefox), що взаємодіє зі Streamlit-сервером через HTTP/WebSocket;',
        'Сервер додатку: Python 3.11+ Runtime із Streamlit Server (порт 8501), '
        'усіма модулями системи, SQLite базою даних та кешем GraphML файлів;',
        'Зовнішні хмарні сервіси: OpenStreetMap (Nominatim + Overpass API) та Mapbox.',
    ]
    for item in deploy_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'Кешування OSM-графів у форматі GraphML дозволяє уникнути повторних '
        'звернень до Overpass API та значно пришвидшує повторні запуски оптимізації.'
    )

    add_image(doc, 'Deployment_DeliveryIQ.png',
              'Рис. 2.10 — Діаграма розгортання системи DeliveryIQ', width=Cm(14))

    doc.add_page_break()

    # =========================================================================
    # РОЗДІЛ 3
    # =========================================================================
    doc.add_heading('3. РЕАЛІЗАЦІЯ', level=1)

    doc.add_heading('3.1. Вибір засобів розробки', level=2)

    doc.add_paragraph('Для реалізації системи обрано наступний технологічний стек:')

    tech = [
        ('Python 3.11+', 'основна мова програмування;'),
        ('Streamlit', 'фреймворк для побудови інтерактивного веб-інтерфейсу;'),
        ('OSMnx 2.x', 'бібліотека для завантаження та аналізу мереж OpenStreetMap;'),
        ('NetworkX', 'бібліотека для роботи з графами (Dijkstra, побудова маршрутів);'),
        ('Folium + AntPath', 'інтерактивна картографічна візуалізація з анімованими маршрутами;'),
        ('geopy (Nominatim)', 'геокодування адрес у координати;'),
        ('scikit-learn (KMeans)', 'географічна кластеризація для VRP;'),
        ('SQLite', 'легка реляційна БД для зберігання пакунків;'),
        ('PlantUML', 'засіб побудови UML-діаграм для проектування.'),
    ]
    for name, desc in tech:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name)
        run.bold = True
        p.add_run(f' — {desc}')

    doc.add_heading('3.2. Структура програмного забезпечення', level=2)

    doc.add_paragraph(
        'Архітектура системи побудована за принципом чіткого розділення '
        'відповідальностей (Separation of Concerns). Конвеєр обробки даних:'
    )

    pipeline = [
        'app.py (Streamlit UI) — точка входу, координація всіх операцій;',
        '→ geocoder.py — перетворення адрес у координати (Nominatim);',
        '→ graph_builder.py — завантаження OSM-мережі, LSCC-обрізка, модальні графи;',
        '→ route_solver.py — матриця відстаней (Dijkstra) + TSP (NN / 2-opt / GA / Christofides);',
        '→ vrp_solver.py — CVRP: розподіл зупинок, K-Means, TSP для кожного ТЗ;',
        '→ visualizer.py — Folium-карта з AntPath-маршрутами;',
        '→ package_db.py — SQLite CRUD для пакунків.',
    ]
    for item in pipeline:
        doc.add_paragraph(item)

    doc.add_heading('3.3. Ключові алгоритми', level=2)

    doc.add_paragraph(
        'Система реалізує декілька алгоритмів оптимізації маршрутів:'
    )

    algorithms = [
        ('Nearest Neighbour (NN)', 'жадібний алгоритм O(n²), використовується для 1-2 зупинок;'),
        ('2-opt', 'ітеративне покращення NN-розв\'язку шляхом обміну ребер, O(n³) за ітерацію;'),
        ('Генетичний алгоритм (GA)', 'Order Crossover (OX) з елітизмом, 120 особин × 400 поколінь, '
         'для задач із >20 зупинками;'),
        ('Christofides', '½-наближення для метричного TSP, потребує повний зв\'язний граф.'),
    ]
    for name, desc in algorithms:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name)
        run.bold = True
        p.add_run(f' — {desc}')

    doc.add_paragraph(
        'Автоматичний вибір алгоритму (method="auto") визначається кількістю '
        'зупинок: n≤2 → NN, 3≤n≤20 → 2-opt, n>20 → генетичний алгоритм. '
        'Це забезпечує баланс між якістю розв\'язку та часом обчислення.'
    )

    doc.add_paragraph(
        'Механізм «останньої милі» для автомобільного режиму: якщо адреса '
        'доставки знаходиться на пішохідній вулиці, автомобіль паркується '
        'на найближчому car-accessible вузлі, а залишок шляху (до 100 м) '
        'долається пішки. Час переїзду включає обидві частини.'
    )

    doc.add_page_break()

    # =========================================================================
    # ВИСНОВКИ
    # =========================================================================
    doc.add_heading('ВИСНОВКИ', level=1)

    doc.add_paragraph(
        'У результаті виконання курсової роботи було спроектовано інформаційну '
        'систему «DeliveryIQ» для оптимізації розподілу вантажів у транспортній '
        'мережі з використанням об\'єктно-орієнтованого підходу.'
    )

    doc.add_paragraph('В ході роботи було:')

    conclusions = [
        'проведено аналіз предметної області маршрутизації транспортних засобів та формалізовано задачу CVRP;',
        'побудовано діаграму прецедентів, що визначає функціональні вимоги системи та її акторів;',
        'розроблено дві діаграми класів: предметної області (ключові сутності та зв\'язки) та модульної структури (розподіл функціональності між модулями);',
        'створено три діаграми послідовності для основних сценаріїв: оптимізація TSP, геокодування адрес, VRP-оптимізація;',
        'побудовано діаграму діяльності процесу оптимізації з паралельними потоками;',
        'розроблено діаграму станів для життєвого циклу пакунку;',
        'побудовано діаграми компонентів та розгортання, що описують фізичну архітектуру системи;',
        'описано технологічний стек та ключові алгоритми оптимізації.',
    ]
    for i, item in enumerate(conclusions, 1):
        doc.add_paragraph(f'{i}) {item}')

    doc.add_paragraph(
        'Загалом побудовано 10 UML-діаграм, що забезпечують повне покриття '
        'статичної структури, динамічної поведінки та фізичної архітектури '
        'системи DeliveryIQ. Розроблений комплект проектної документації '
        'може бути використаний як основа для подальшої реалізації, '
        'тестування та супроводу системи.'
    )

    doc.add_page_break()

    # =========================================================================
    # СПИСОК ДЖЕРЕЛ
    # =========================================================================
    doc.add_heading('СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ', level=1)

    sources = [
        'Boeing G. OSMnx: New methods for acquiring, constructing, analyzing, and visualizing complex street networks. Computers, Environment and Urban Systems, 2017. Vol. 65. P. 126-139.',
        'Toth P., Vigo D. The Vehicle Routing Problem. SIAM Monographs on Discrete Mathematics and Applications, 2002.',
        'Christofides N. Worst-case analysis of a new heuristic for the travelling salesman problem. Report 388, Graduate School of Industrial Administration, Carnegie Mellon University, 1976.',
        'Goldberg D.E. Genetic Algorithms in Search, Optimization, and Machine Learning. Addison-Wesley, 1989.',
        'Fowler M. UML Distilled: A Brief Guide to the Standard Object Modeling Language. 3rd ed. Addison-Wesley, 2004.',
        'Booch G., Rumbaugh J., Jacobson I. The Unified Modeling Language User Guide. 2nd ed. Addison-Wesley, 2005.',
        'Streamlit documentation. URL: https://docs.streamlit.io/',
        'NetworkX documentation. URL: https://networkx.org/documentation/',
        'Folium documentation. URL: https://python-visualization.github.io/folium/',
        'OpenStreetMap Wiki. URL: https://wiki.openstreetmap.org/',
    ]
    for i, src in enumerate(sources, 1):
        doc.add_paragraph(f'{i}. {src}')

    # Save
    doc.save(str(OUTPUT))
    print("Document saved successfully.")


if __name__ == '__main__':
    build_document()
